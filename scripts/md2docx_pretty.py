#!/usr/bin/env python3
"""Render the PRD markdown to a polished .docx (python-docx).

Fixes over the generic converter: joins hard-wrapped lines into single
paragraphs, joins multi-line blockquotes into one callout, adds a cover
block + metadata table, a TOC field, colored heading styles, shaded table
headers with fitted column widths, and footer page numbers.
"""
import io, os, re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import sys
if len(sys.argv) != 3:
    sys.exit("Usage: python scripts/md2docx_pretty.py <in.md> <out.docx>")
SRC, DST = sys.argv[1], sys.argv[2]

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
BLUE = RGBColor(0x2E, 0x5A, 0x88)
GREY = RGBColor(0x44, 0x44, 0x44)
HDR_FILL = "1F3A5F"
ZEBRA_FILL = "EEF2F7"
CALLOUT_FILL = "FBF3D5"
INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")
USABLE_CM = 16.5  # A4 (21cm) minus 2x2.25cm margins


def add_inline(paragraph, text):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"; r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def shade(cell, fill):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(sh)


def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}"); e.set(qn("w:w"), str(val)); e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def page_number_footer(section):
    p = section.footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SIDESA-CM — PRD v0.3      ·      Halaman ")
    run.font.size = Pt(8); run.font.color.rgb = GREY
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    for tag, attrs in (("w:fldChar", {"w:fldCharType": "begin"}),):
        e = OxmlElement(tag)
        for k, v in attrs.items():
            e.set(qn(k), v)
        run._r.append(e)
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    run._r.append(instr)
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate"); run._r.append(sep)
    hint = OxmlElement("w:t"); hint.text = "Klik kanan → Update Field untuk memperbarui daftar isi."
    run._r.append(hint)
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end"); run._r.append(end)


def style_doc(doc):
    n = doc.styles["Normal"]
    n.font.name = "Calibri"; n.font.size = Pt(10.5)
    n.paragraph_format.space_after = Pt(6); n.paragraph_format.line_spacing = 1.15
    for name, size, color, before in (("Heading 1", 14, NAVY, 14), ("Heading 2", 11.5, BLUE, 10)):
        s = doc.styles[name]
        s.font.name = "Calibri"; s.font.size = Pt(size); s.font.bold = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(before); s.paragraph_format.space_after = Pt(4)
        s.paragraph_format.keep_with_next = True
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21)
    for m in ("top", "bottom", "left", "right"):
        setattr(sec, f"{m}_margin", Cm(2.25))


def col_widths(rows):
    ncol = len(rows[0])
    avg = [0] * ncol
    for r in rows:
        for i in range(ncol):
            avg[i] = max(avg[i], len(re.sub(r"[*`]", "", r[i])) if i < len(r) else 0)
    tot = sum(avg) or ncol
    widths = [max(1.6, USABLE_CM * a / tot) for a in avg]
    scale = USABLE_CM / sum(widths)
    return [w * scale for w in widths]


def add_table(doc, rows):
    header, body = rows[0], rows[1:]
    widths = col_widths(rows)
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; t.allow_autofit = False
    for cell, text, w in zip(t.rows[0].cells, header, widths):
        cell.width = Cm(w); shade(cell, HDR_FILL); set_cell_margins(cell)
        cell.text = ""; add_inline(cell.paragraphs[0], text)
        for run in cell.paragraphs[0].runs:
            run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); run.font.size = Pt(10)
    for ri, row in enumerate(body):
        cells = t.add_row().cells
        for i, (cell, text) in enumerate(zip(cells, row + [""] * (len(header) - len(row)))):
            cell.width = Cm(widths[i]); set_cell_margins(cell)
            if ri % 2 == 1:
                shade(cell, ZEBRA_FILL)
            cell.text = ""; add_inline(cell.paragraphs[0], text)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9.5)
    doc.add_paragraph()


def add_callout(doc, text):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; t.allow_autofit = False
    cell = t.rows[0].cells[0]; cell.width = Cm(USABLE_CM)
    shade(cell, CALLOUT_FILL); set_cell_margins(cell, 100, 100, 160, 160)
    cell.text = ""; add_inline(cell.paragraphs[0], text)
    doc.add_paragraph()


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_sep(line):
    return bool(re.fullmatch(r"\|[\s:|-]+\|", line.strip()))


LIST_RE = re.compile(r"^([-*]|\d+\.) ")


def blocks(lines):
    """Group into blocks separated by blank lines."""
    buf, out = [], []
    for ln in lines:
        if ln.strip():
            buf.append(ln)
        elif buf:
            out.append(buf); buf = []
    if buf:
        out.append(buf)
    return out


def render_list(doc, block):
    items = []
    for ln in block:
        s = ln.strip()
        if LIST_RE.match(s):
            items.append([s])
        elif items:
            items[-1].append(s)  # wrapped continuation joins the item
        else:
            items.append([s])
    for it in items:
        s = " ".join(it)
        ordered = bool(re.match(r"^\d+\. ", s))
        body = re.sub(r"^([-*]|\d+\.) ", "", s)
        p = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
        add_inline(p, body)


def main():
    raw = io.open(SRC, encoding="utf-8").read()
    # normalise inline math to unicode
    raw = re.sub(r"\\text\{([^}]*)\}", r"\1", raw).replace(r"\cdot", "·").replace(r"\|", "‖")
    raw = re.sub(r"\$([^$]*)\$", r"\1", raw)
    lines = raw.splitlines()

    doc = Document()
    style_doc(doc)
    seen_heading = False

    for block in blocks(lines):
        first = block[0].strip()

        if first.startswith("|"):
            rows = [split_row(l) for l in block if l.strip().startswith("|") and not is_sep(l)]
            if rows:
                add_table(doc, rows)
            continue

        if first == "---":
            continue

        if first.startswith("# ") and not seen_heading:
            # cover title
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(first[2:]); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY
            p.paragraph_format.space_after = Pt(4)
            continue

        if first.startswith("## "):
            seen_heading and None
            if not seen_heading:
                seen_heading = True
            doc.add_heading(first[3:], level=1)
            continue
        if first.startswith("### "):
            doc.add_heading(first[3:], level=2)
            continue

        if first.startswith("> "):
            add_callout(doc, " ".join(l.strip()[2:] for l in block))
            continue

        # metadata block on the cover: every line **Key:** value
        if not seen_heading and all(l.strip().startswith("**") and ":**" in l for l in block):
            t = doc.add_table(rows=0, cols=2)
            t.autofit = False; t.allow_autofit = False
            for l in block:
                key, val = l.strip().split(":**", 1)
                key = key.lstrip("*").strip()
                cells = t.add_row().cells
                cells[0].width = Cm(3.2); cells[1].width = Cm(USABLE_CM - 3.2)
                cells[0].text = ""; kr = cells[0].paragraphs[0].add_run(key + ":")
                kr.bold = True; kr.font.color.rgb = BLUE
                cells[1].text = ""; add_inline(cells[1].paragraphs[0], val.strip())
            doc.add_paragraph()
            # TOC after the metadata/cover, on its own page
            hp = doc.add_paragraph(); hr = hp.add_run("Daftar Isi")
            hr.bold = True; hr.font.size = Pt(13); hr.font.color.rgb = NAVY
            add_toc(doc)
            continue

        # ordinary list?
        if LIST_RE.match(first):
            render_list(doc, block)
            continue

        # body paragraph: join wrapped lines
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(p, " ".join(l.strip() for l in block))

    page_number_footer(doc.sections[0])
    doc.save(DST)

    d = Document(DST)
    heads = sum(1 for p in d.paragraphs if p.style.name.startswith(("Heading", "Title")))
    print(f"OK {DST}")
    print(f"   {os.path.getsize(DST)} B | paragraf={len(d.paragraphs)} heading={heads} "
          f"tabel={len(d.tables)} sisa-$={sum(p.text.count('$') for p in d.paragraphs)}")


if __name__ == "__main__":
    main()
