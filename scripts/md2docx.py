#!/usr/bin/env python3
"""Render a project Markdown doc to .docx.

Handles the subset of Markdown these docs actually use: headings, paragraphs,
bullet lists, pipe tables, blockquotes, horizontal rules, and inline
**bold** / `code` / *italic*.

Usage:  python scripts/md2docx.py docs/LAPORAN.md docs/LAPORAN.docx
"""
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")


def add_inline(paragraph, text):
    """Split text into runs so **bold**, `code` and *italic* render properly."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_separator(line):
    return bool(re.fullmatch(r"\|[\s:|-]+\|", line.strip()))


def add_table(doc, rows):
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, header):
        cell.text = ""
        add_inline(cell.paragraphs[0], text)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in body:
        cells = table.add_row().cells
        # Tolerate ragged rows rather than crashing on a stray pipe.
        for cell, text in zip(cells, row + [""] * (len(header) - len(row))):
            cell.text = ""
            add_inline(cell.paragraphs[0], text)
    doc.add_paragraph()


def convert(md_path, docx_path):
    lines = open(md_path, encoding="utf-8").read().splitlines()
    doc = Document()
    table_buf = []
    i = 0

    def flush_table():
        nonlocal table_buf
        if table_buf:
            add_table(doc, table_buf)
            table_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # --- pipe table -------------------------------------------------
        if stripped.startswith("|") and stripped.endswith("|"):
            if not is_separator(stripped):
                table_buf.append(split_row(stripped))
            i += 1
            continue
        flush_table()

        if not stripped or stripped == "---":
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith("> "):
            para = doc.add_paragraph(style="Intense Quote")
            add_inline(para, stripped[2:])
        elif re.match(r"^[-*] ", stripped):
            para = doc.add_paragraph(style="List Bullet")
            add_inline(para, stripped[2:])
        elif re.match(r"^\d+\. ", stripped):
            para = doc.add_paragraph(style="List Number")
            add_inline(para, re.sub(r"^\d+\. ", "", stripped))
        else:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline(para, stripped)
        i += 1

    flush_table()
    doc.save(docx_path)
    print(f"OK  {md_path} -> {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        sys.exit(__doc__)
    convert(sys.argv[1], sys.argv[2])
